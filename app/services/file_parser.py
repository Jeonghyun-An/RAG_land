# app/services/file_parser.py
from __future__ import annotations
from typing import Dict, List, Tuple, Union, Optional
import os, re
from io import BytesIO
from app.services.ocr_service import _norm_easyocr_langs

"""
PDF 텍스트 추출
- OCR 모드: OCR_MODE in {"auto","always","never"} (default=auto)
- OCR 엔진: OCR_ENGINE in {"paddle","tesseract","easyocr"} (default=paddle)
- 언어:
    * paddle : "korean" 또는 "en", "ch" 등
    * tesseract : "kor+eng" 같은 조합
    * easyocr : "ko,en" 같은 콤마구분
- 렌더링 DPI: OCR_DPI (default=300)
- 추가 옵션:
    * OCR_TESSERACT_CMD: pytesseract 실행 파일 경로(Windows 등)
    * OCR_EASYOCR_GPU: "1"이면 GPU 사용 시도, 기본 "0"
    * OCR_MIN_CHARS: auto 모드에서 OCR 전환 기준(기본 40)
"""
from PIL import Image
if not hasattr(Image, "ANTIALIAS"):
    try:
        Image.ANTIALIAS = Image.LANCZOS
    except Exception:
        Image.ANTIALIAS = getattr(Image, "BICUBIC", None)

def _pdf_page_count(path: str) -> int:
    """PDF 페이지 수 안전 획득"""
    try:
        import fitz
        with fitz.open(path) as d:
            return d.page_count
    except Exception:
        return 1  # 최후 보호

def _make_image_placeholder_chunk(page_no: int) -> tuple[str, dict]:
    text = f"[page {page_no}: image or low-text content]"
    meta = {
        "type": "image_page",
        "section": "",
        "page": page_no,
        "pages": [page_no],
        "token_count": len(text.split()),
        "bboxes": {},
    }
    return (text, meta)

# ---------------------- Text extract (no OCR) ---------------------- #
def _extract_text_pdfminer(path: str, by_page: bool) -> Union[str, List[Tuple[int, str]]]:
    import pdfminer.high_level
    from pdfminer.layout import LAParams, LTTextContainer

    laparams = LAParams()
    if not by_page:
        return (pdfminer.high_level.extract_text(path, laparams=laparams) or "").strip()

    pages: List[Tuple[int, str]] = []
    for i, layout in enumerate(pdfminer.high_level.extract_pages(path, laparams=laparams), start=1):
        parts = []
        for elem in layout:
            if isinstance(elem, LTTextContainer):
                parts.append(elem.get_text())
        text = "".join(parts).strip()
        if text:
            pages.append((i, text))
    return pages


def _extract_text_pypdf2(path: str, by_page: bool) -> Union[str, List[Tuple[int, str]]]:
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    if not by_page:
        txt = "\n\n".join([(p.extract_text() or "") for p in reader.pages]).strip()
        return txt
    pages: List[Tuple[int, str]] = []
    for i, p in enumerate(reader.pages, start=1):
        t = (p.extract_text() or "").strip()
        if t:
            pages.append((i, t))
    return pages

# ---------------------- Common rendering (PyMuPDF) ---------------------- #
def _render_pdf_pages_fitz(path: str, dpi: int):
    import fitz
    import numpy as np

    zoom = max(1.0, dpi / 72.0)
    mat = fitz.Matrix(zoom, zoom)
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            yield i, img
    finally:
        doc.close()


# ---------------------- OCR engines (via fitz images) ---------------------- #
def _ocr_with_paddle(images_iter, by_page: bool, lang: str) -> Union[str, List[Tuple[int, str]]]:
    from paddleocr import PaddleOCR
    ocr = PaddleOCR(lang=("korean" if lang in ("kor", "korean", "ko") else lang), show_log=False)
    pages: List[Tuple[int, str]] = []
    for pno, img in images_iter:
        result = ocr.ocr(img, cls=True)
        text = " ".join([line[1][0] for line in (result[0] or []) if line[1] and line[1][0]]).strip()
        if text:
            pages.append((pno, text))
    if by_page:
        return pages
    return "\n\n".join([t for _, t in pages]).strip()

def _ocr_with_tesseract(images_iter, by_page: bool, lang: str) -> Union[str, List[Tuple[int, str]]]:
    import pytesseract
    from PIL import Image
    
    tcmd = os.getenv("OCR_TESSERACT_CMD")
    if tcmd:
        pytesseract.pytesseract.tesseract_cmd = tcmd

    pages: List[Tuple[int, str]] = []
    for pno, img in images_iter:
        pil = Image.fromarray(img)
        text = pytesseract.image_to_string(pil, lang=(lang or "kor+eng")).strip()
        if text:
            pages.append((pno, text))
    if by_page:
        return pages
    return "\n\n".join([t for _, t in pages]).strip()

def _ocr_with_easyocr(images_iter, by_page: bool, lang: str) -> Union[str, List[Tuple[int, str]]]:
    """
    EasyOCR로 이미지에서 텍스트 추출
    - CID 코드 자동 감지
    - 회전된 이미지 자동 보정
    """
    import easyocr
    import numpy as np
    from PIL import Image
    
    langs = _norm_easyocr_langs(lang)
    gpu = os.getenv("OCR_EASYOCR_GPU", "1").strip() == "1"
    
    # 옵션
    width_ths = float(os.getenv("OCR_WIDTH_THS", "0.5"))
    height_ths = float(os.getenv("OCR_HEIGHT_THS", "0.5"))
    auto_rotate = os.getenv("OCR_AUTO_ROTATE", "1").strip() == "1"
    
    reader = easyocr.Reader(langs, gpu=gpu)

    pages: List[Tuple[int, str]] = []
    for pno, img in images_iter:
        # PIL Image → numpy array
        if hasattr(img, 'mode'):
            img_pil = img
            img_array = np.array(img)
        else:
            img_array = img
            img_pil = Image.fromarray(img_array)
        
        best_text = ""
        best_conf = 0
        
        # ★ 회전 각도별 시도 (auto_rotate=True 시)
        if auto_rotate:
            angles = [0, 90, 180, 270]
            for angle in angles:
                # 이미지 회전
                if angle == 0:
                    test_img = img_array
                else:
                    rotated_pil = img_pil.rotate(angle, expand=True)
                    test_img = np.array(rotated_pil)
                
                # OCR 실행
                try:
                    res = reader.readtext(
                        test_img,
                        detail=1,
                        paragraph=False,
                        width_ths=width_ths,
                        height_ths=height_ths,
                    )
                except Exception as e:
                    print(f"⚠️ OCR failed at {angle}°: {e}")
                    continue
                
                if not res:
                    continue
                
                # 신뢰도 계산
                total_conf = sum([x[2] for x in res if len(x) >= 3])
                text_count = len([x for x in res if len(x) >= 2 and len(x[1].strip()) > 1])
                avg_conf = total_conf / len(res) if res else 0
                
                # 점수 = 평균 신뢰도 × 텍스트 수
                score = avg_conf * text_count
                
                # 텍스트 추출
                sorted_res = sorted(res, key=lambda x: (x[0][0][1], x[0][0][0]))
                text_lines = [x[1] for x in sorted_res if len(x) >= 2 and x[1]]
                text = " ".join(text_lines).strip()
                
                # 최고 점수 갱신
                if score > best_conf:
                    best_conf = score
                    best_text = text
                    if angle != 0:
                        print(f"✅ Page {pno}: Best orientation at {angle}° (score: {score:.2f})")
        else:
            # 회전 감지 비활성화 시 바로 실행
            res = reader.readtext(
                img_array,
                detail=1,
                paragraph=False,
                width_ths=width_ths,
                height_ths=height_ths,
            )
            
            if res:
                sorted_res = sorted(res, key=lambda x: (x[0][0][1], x[0][0][0]))
                best_text = " ".join([x[1] for x in sorted_res if len(x) >= 2 and x[1]]).strip()
        
        # 결과 저장
        if best_text:
            pages.append((pno, best_text))
    
    if by_page:
        return pages
    return "\n\n".join([t for _, t in pages]).strip()

def _has_cid_codes(text: str) -> bool:
    """
    텍스트에 CID 코드가 많으면 True 반환
    
    CID 코드 예시: (cid:31), (cid:144), (cid:157) 등
    """
    if not text:
        return False
    
    # (cid:숫자) 패턴 찾기
    cid_pattern = r'\(cid:\d+\)'
    cid_matches = re.findall(cid_pattern, text)
    cid_count = len(cid_matches)
    
    # 임계값: 10개 이상
    if cid_count < 10:
        return False
    
    # 전체 텍스트 대비 CID 비율
    total_chars = len(text)
    cid_chars = sum(len(m) for m in cid_matches)
    cid_ratio = cid_chars / max(1, total_chars)
    
    # CID가 5% 이상이면 비정상
    is_cid = cid_ratio > 0.05
    
    if is_cid:
        print(f"⚠️ CID codes detected: {cid_count} codes, {cid_ratio*100:.1f}% of text")
    
    return is_cid

# ---------------------- Heuristic for OCR fallback ---------------------- #
def _should_ocr(txt_or_pages: Union[str, List[Tuple[int, str]]]) -> bool:
    """텍스트 밀도 낮거나 CID 코드 많으면 OCR로 전환"""
    try:
        th = int(os.getenv("OCR_MIN_CHARS", "40"))
    except Exception:
        th = 40
    
    if isinstance(txt_or_pages, str):
        if len(txt_or_pages.strip()) < th:
            return True
        if _has_cid_codes(txt_or_pages):  # ← 추가
            return True
        return False
    
    total = sum(len(t or "") for _, t in (txt_or_pages or []))
    if total < th:
        return True
    
    # 각 페이지에서 CID 코드 체크
    for _, text in (txt_or_pages or []):
        if _has_cid_codes(text):  # ← 추가
            return True
    
    return False

# ---------------------- Public API ---------------------- #
def parse_docx_sections(path: str) -> List[Tuple[int, str]]:
    """DOCX를 섹션 단위로 반환"""
    from docx import Document
    doc = Document(path)
    sections: List[Tuple[int, str]] = []
    cur = []
    sec_no = 0

    def flush():
        nonlocal cur, sec_no
        if cur:
            sec_no += 1
            sections.append((sec_no, "\n".join(cur).strip()))
            cur = []

    for p in doc.paragraphs:
        style = (p.style.name or "").lower() if p.style else ""
        line = p.text.strip()
        if not line:
            continue
        if "heading" in style:
            flush()
            cur.append(line)
        else:
            cur.append(line)

    for t in doc.tables:
        lines = []
        for r in t.rows:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            lines.append(" | ".join(cells))
        if lines:
            cur.append("[표]\n" + "\n".join(lines))

    flush()
    return sections

def parse_xlsx_tables(path: str) -> List[Tuple[int, str]]:
    """XLSX를 시트 단위로 반환"""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    out: List[Tuple[int,str]] = []
    for si, ws in enumerate(wb.worksheets, start=1):
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h or "").strip() for h in rows[0]]
        lines = [f"[시트] {ws.title}"]
        for r in rows[1:]:
            pairs = []
            for h, v in zip(headers, r):
                if h:
                    pairs.append(f"{h}: {'' if v is None else str(v)}")
            if pairs:
                lines.append(", ".join(pairs))
        txt = "\n".join(lines).strip()
        if txt:
            out.append((si, txt))
    return out

def parse_pdf(path: str, by_page: bool = False) -> Union[str, List[Tuple[int, str]]]:
    """PDF 텍스트/OCR 추출"""
    ocr_mode   = os.getenv("OCR_MODE", "auto").lower()
    ocr_engine = os.getenv("OCR_ENGINE", "easyocr").lower()
    ocr_dpi    = int(os.getenv("OCR_DPI", "300"))

    #  페이지 수를 미리 구해둔다 (폴백 생성용)
    page_count = _pdf_page_count(path)

    if ocr_engine == "tesseract":
        ocr_langs = os.getenv("OCR_LANGS", "kor+eng")
    elif ocr_engine == "easyocr":
        ocr_lang = os.getenv("OCR_LANG", "ko,en")
    else:
        ocr_lang = os.getenv("OCR_LANG", "korean")

    text_result: Optional[Union[str, List[Tuple[int, str]]]] = None
    if ocr_mode in ("auto", "never"):
        try:
            text_result = _extract_text_pdfminer(path, by_page)
        except Exception:
            try:
                text_result = _extract_text_pypdf2(path, by_page)
            except Exception:
                text_result = "" if not by_page else []
        if ocr_mode == "never":
            return text_result if text_result is not None else ("" if not by_page else [])
        if text_result and not _should_ocr(text_result):
            return text_result

    # OCR 시도 (1차 → 2차 대체 엔진)
    try:
        images_iter = _render_pdf_pages_fitz(path, ocr_dpi)

        def _empty(x):
            return (isinstance(x, str) and not x.strip()) or (isinstance(x, list) and len(x) == 0)

        # 1차: 지정 엔진
        if ocr_engine == "tesseract":
            ocr_out = _ocr_with_tesseract(images_iter, by_page, ocr_langs)
        elif ocr_engine == "easyocr":
            ocr_out = _ocr_with_easyocr(images_iter, by_page, ocr_lang)
        else:
            ocr_out = _ocr_with_paddle(images_iter, by_page, ocr_lang)

        # 1차가 비었으면 대체 엔진 재시도
        if _empty(ocr_out):
            images_iter = _render_pdf_pages_fitz(path, ocr_dpi)  # 제너레이터 재생성
            if ocr_engine == "easyocr":
                # easyocr → tesseract
                ocr_out = _ocr_with_tesseract(images_iter, by_page, os.getenv("OCR_LANGS", "kor+eng"))
            else:
                # tesseract/paddle → easyocr
                ocr_out = _ocr_with_easyocr(images_iter, by_page, os.getenv("OCR_LANG", "ko,en"))

        # 그래도 비면: 텍스트 추출값이 있으면 그걸, 없으면 페이지별 플레이스홀더
        if _empty(ocr_out):
            if text_result and ((isinstance(text_result, str) and text_result.strip()) or (isinstance(text_result, list) and len(text_result) > 0)):
                return text_result
            # 페이지별 플레이스홀더 생성
            if by_page:
                return [(p, _make_image_placeholder_chunk(p)[0]) for p in range(1, page_count + 1)]
            else:
                return "\n".join(_make_image_placeholder_chunk(p)[0] for p in range(1, page_count + 1))

        return ocr_out

    except Exception as e:
        # OCR 중 예외: 텍스트 추출값이 있으면 사용, 아니면 페이지별 플레이스홀더
        if text_result and ((isinstance(text_result, str) and text_result.strip()) or (isinstance(text_result, list) and len(text_result) > 0)):
            return text_result
        if by_page:
            return [(p, _make_image_placeholder_chunk(p)[0]) for p in range(1, page_count + 1)]
        else:
            return "\n".join(_make_image_placeholder_chunk(p)[0] for p in range(1, page_count + 1))



def parse_pdf_blocks(path: str) -> list[tuple[int, list[dict]]]:
    import fitz
    out = []
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc, start=1):
            blocks = []
            for b in page.get_text("blocks"):
                if not b or len(b) < 5:
                    continue
                x0, y0, x1, y1, txt = b[0], b[1], b[2], b[3], (b[4] or "").strip()
                if txt:
                    blocks.append({
                        "text": txt,
                        "bbox": {"x0": float(x0), "y0": float(y0), "x1": float(x1), "y1": float(y1)}
                    })
            out.append((i, blocks))
    finally:
        doc.close()
    return out


    
def parse_any(path: str) -> List[Tuple[int, str]]:
    """파일 확장자에 따라 적절한 파서 선택"""
    ext = (os.path.splitext(path)[1] or "").lower()
    direct_docx = os.getenv("RAG_PARSE_DIRECT_DOCX", "1") == "1"
    direct_xlsx = os.getenv("RAG_PARSE_DIRECT_XLSX", "1") == "1"
    allow_convert = os.getenv("RAG_CONVERT_NONPDF_TO_PDF", "1") == "1"

    if ext == ".pdf":
        return parse_pdf(path, by_page=True)

    if ext in (".docx",) and direct_docx:
        return parse_docx_sections(path)

    if ext in (".xlsx", ".xlsm", ".csv") and direct_xlsx:
        if ext == ".csv":
            import csv
            lines = []
            with open(path, "r", encoding="utf-8", errors="ignore",newline="") as f:
                rdr = csv.reader(f)
                for row in rdr:
                    lines.append(" | ".join([c.strip() for c in row]))
            return [(1, "\n".join(lines))]
        return parse_xlsx_tables(path)

    if allow_convert:
        from app.services.pdf_converter import convert_to_pdf
        pdf_path = convert_to_pdf(path)
        return parse_pdf(pdf_path, by_page=True)

    raise RuntimeError(f"Unsupported file type: {ext}")

def sniff_ext_from_name(name: str) -> str:
    return (os.path.splitext(name)[1] or "").lower()

def parse_pdf_pages_from_bytes(pdf_bytes: bytes) -> List[str]:
    """bytes에서 페이지별 텍스트 추출 (pdfminer.six 기반)"""
    import pdfminer.high_level
    from pdfminer.layout import LAParams, LTTextContainer
    from io import BytesIO
    
    laparams = LAParams()
    pages: List[str] = []
    for layout in pdfminer.high_level.extract_pages(BytesIO(pdf_bytes), laparams=laparams):
        parts = []
        for elem in layout:
            if isinstance(elem, LTTextContainer):
                parts.append(elem.get_text())
        text = "".join(parts).strip()
        pages.append(text)
    return pages


def parse_pdf_blocks_from_bytes(pdf_bytes: bytes) -> List[Tuple[int, List[Dict]]]:
    """bytes에서 레이아웃 블록 추출"""
    import fitz
    out = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        for i, page in enumerate(doc, start=1):
            blocks = []
            for b in page.get_text("blocks"):
                if not b or len(b) < 5:
                    continue
                x0, y0, x1, y1, txt = b[0], b[1], b[2], b[3], (b[4] or "").strip()
                if txt:
                    blocks.append({"text": txt, "bbox": [float(x0), float(y0), float(x1), float(y1)]})
            out.append((i, blocks))
    finally:
        doc.close()
    return out


def parse_docx_from_bytes(content: bytes) -> List[Dict]:
    """DOCX bytes → 라인 단위"""
    from docx import Document
    from io import BytesIO
    
    doc = Document(BytesIO(content))
    items = []
    for p in doc.paragraphs:
        line = p.text.strip()
        if line:
            items.append({"text": line})
    return items


def parse_plaintext_bytes(content: bytes) -> List[Dict]:
    """평문 bytes → 라인 단위"""
    try:
        text = content.decode("utf-8", errors="ignore")
    except:
        text = ""
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    return [{"text": ln} for ln in lines]


def parse_any_bytes(name_hint: str, content: bytes) -> dict:
    """
    bytes에서 파일 형식 감지 후 파싱
    반환 형식:
      {
        "kind": "pdf"|"docx"|"plain",
        "ext": ".pdf"|...,
        "pages": [...],     # PDF면 페이지별 텍스트
        "blocks": [...],    # PDF면 레이아웃 블록
        "items": [{"text":..}, ...],  # DOCX/PLAIN 등
      }
    """
    ext = sniff_ext_from_name(name_hint)

    if ext == ".pdf":
        pages = parse_pdf_pages_from_bytes(content)
        blocks = parse_pdf_blocks_from_bytes(content)
        return {"kind": "pdf", "ext": ext, "pages": pages, "blocks": blocks}

    if ext == ".docx":
        items = parse_docx_from_bytes(content)
        return {"kind": "docx", "ext": ext, "items": items}

    if ext in (".hwpx", ".hwp"):
        from app.services.pdf_converter import convert_stream_to_pdf_bytes
        pdf_bytes = convert_stream_to_pdf_bytes(content, ext)
        if pdf_bytes:
            pages = parse_pdf_pages_from_bytes(pdf_bytes)
            blocks = parse_pdf_blocks_from_bytes(pdf_bytes)
            return {"kind": "pdf", "ext": ".pdf", "pages": pages, "blocks": blocks}
        return {"kind": "plain", "ext": ext, "items": parse_plaintext_bytes(content)}

    return {"kind": "plain", "ext": ext, "items": parse_plaintext_bytes(content)}


